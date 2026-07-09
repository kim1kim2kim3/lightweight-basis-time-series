from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = Path("Ours_국내학회_수정본_수식포함.docx")


ACCURACY_ROWS = [
    ["96", "0.2310", "0.3031", "0.2119", "0.2794", "0.2239", "0.2900"],
    ["192", "0.2514", "0.3275", "0.2396", "0.3133", "0.2489", "0.3201"],
    ["336", "0.2655", "0.3393", "0.2553", "0.3264", "0.2589", "0.3283"],
    ["720", "0.3221", "0.3987", "0.2879", "0.3673", "0.2910", "0.3667"],
]

PARAM_ROWS = [
    ["96", "18.6K", "6903.9K", "11.1K"],
    ["192", "37.2K", "7493.8K", "11.1K"],
    ["336", "65.2K", "8378.7K", "11.1K"],
    ["720", "139.7K", "10738.4K", "11.1K"],
]


def set_run_font(run, *, size: float, bold: bool = False, ascii_font: str = "Times New Roman", east_asia_font: str = "바탕") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def add_paragraph(doc: Document, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size: float = 10, bold: bool = False, space_before=0, space_after=0, first_line_indent_cm: float | None = 0.6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if first_line_indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc: Document, text: str) -> None:
    add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=11,
        bold=True,
        space_before=8,
        space_after=4,
        first_line_indent_cm=None,
    )


def add_equation(doc: Document, text: str) -> None:
    add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10,
        bold=False,
        space_before=2,
        space_after=2,
        first_line_indent_cm=None,
    )


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def fill_table(table, rows: list[list[str]], *, font_size: float = 9) -> None:
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx].cells
        for c_idx, text in enumerate(row):
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_run_font(run, size=font_size, bold=(r_idx == 0))


def add_table_caption(doc: Document, text: str) -> None:
    add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=9.5,
        bold=False,
        space_before=4,
        space_after=2,
        first_line_indent_cm=None,
    )


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Ours: 장기 시계열 예측을 위한 파라미터 효율적 구조 합성 모델")
    set_run_font(title_run, size=16, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Ours: A Parameter-Efficient Structural Synthesis Model for Long-Term Time Series Forecasting")
    set_run_font(subtitle_run, size=11, bold=True)

    for text, size in [
        ("[저자명1], [저자명2]*", 11),
        ("[소속기관명], *[공동소속기관명]", 10),
        ("author1@domain.ac.kr, *author2@domain.ac.kr", 10),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=size, bold=False)

    add_heading(doc, "요 약")
    add_paragraph(
        doc,
        "장기 시계열 예측에서는 horizon이 길어질수록 예측 오차뿐 아니라 모델 규모와 계산 비용도 함께 증가한다. "
        "본 논문은 미래 구간 전체를 큰 출력층으로 직접 회귀하는 대신, 과거 시계열의 잠재 요약으로부터 구조적 basis를 생성하고 "
        "이를 조합하여 horizon 전체를 합성하는 Ours를 제안한다. ETTh1 데이터셋에서 입력 길이 96, 예측 길이 96, 192, 336, 720 조건으로 "
        "Ours를 DLinear와 PatchTST와 비교한 결과, Ours는 모든 horizon에서 11,149개의 파라미터를 유지하면서 DLinear보다 낮은 "
        "scaled MAE를 기록하였다. 또한 336-step과 720-step에서는 PatchTST의 최고 scaled MAE 대비 2% 이내의 오차를 유지하면서 "
        "훨씬 작은 모델 규모를 보였다. 이 결과는 Ours가 장기 시계열 예측에서 정확도와 파라미터 효율성 사이의 실용적인 절충점을 제공함을 보여준다.",
        size=10,
        space_after=2,
    )

    add_heading(doc, "Ⅰ. 서 론")
    add_paragraph(
        doc,
        "장기 시계열 예측은 전력 수요, 설비 상태, 교통 흐름 등 다양한 응용에서 중요한 문제이며, 예측 horizon이 길어질수록 오차 누적과 "
        "장기 의존성 모델링의 난도가 함께 증가한다. 특히 미래 구간이 길어질수록 출력층 규모와 계산 비용까지 커지기 때문에, 정확도뿐 아니라 "
        "구조적 효율성도 함께 고려한 모델 설계가 필요하다.",
    )
    add_paragraph(
        doc,
        "최근 Informer[1], Autoformer[2], FEDformer[3]와 같은 Transformer 계열 모델과 PatchTST[4], TimesNet[5], "
        "iTransformer[8] 등이 장기 예측 성능을 개선해 왔다. 반면 Are Transformers Effective for Time Series Forecasting?[6]은 "
        "단순한 linear 계열 모델도 강한 기준선이 될 수 있음을 보였으며, 이는 backbone의 복잡성만으로는 장기 예측 성능을 충분히 설명하기 어렵다는 점을 시사한다.",
    )
    add_paragraph(
        doc,
        "본 논문은 이러한 문제의식 아래 Ours를 제안한다. Ours는 미래값 전체를 큰 출력층으로 직접 예측하는 대신, 과거 시계열의 잠재 요약으로부터 "
        "구조적 basis를 생성하고 이를 조합하여 horizon 전체를 합성한다. 본 연구의 초점은 모든 조건에서 최고 정확도를 달성하는 데 있지 않으며, "
        "적은 파라미터로도 실용적인 accuracy-efficiency trade-off를 확보할 수 있는지를 검증하는 데 있다.",
    )
    add_paragraph(
        doc,
        "본 논문의 기여는 다음과 같다. 첫째, 미래 구간을 구조적 basis 조합으로 생성하는 Ours를 제안한다. 둘째, horizon 증가와 무관하게 모델 "
        "파라미터 수를 일정하게 유지하는 설계를 제시한다. 셋째, ETTh1 실험을 통해 장기 예측에서의 accuracy-efficiency trade-off를 실증적으로 확인한다.",
    )

    add_heading(doc, "Ⅱ. 제안 방법")
    add_paragraph(
        doc,
        "입력 다변량 시계열을 X ∈ R^(L×C), 예측 horizon을 H라 하자. Ours는 입력 시계열을 causal encoder에 통과시켜 시간 순서를 보존하는 "
        "잠재 표현과 요약 벡터를 얻는다.",
    )
    add_equation(doc, "E = f_enc(X),      z = Pool(E)")
    add_paragraph(
        doc,
        "그 다음 요약 벡터 z로부터 미래 구간을 설명하는 구조적 basis와 각 basis의 계수를 예측한다. 본 연구에서 사용하는 basis는 추세형(trend), "
        "주기형(seasonal), 과도형(transient) 성분으로 구성되며, Ours의 핵심은 horizon의 각 시점을 독립적으로 예측하지 않고 horizon 전체를 "
        "하나의 구조적 조합 문제로 다룬다는 점이다.",
    )
    add_equation(doc, "B = g_basis(z),      a = g_coef(z),      y_hat = Σ_{k=1}^{K} a_k B_k")
    add_paragraph(
        doc,
        "trend basis는 장기 증가·감소 경향을, seasonal basis는 반복 주기를, transient basis는 단기 감쇠 패턴을 담당한다. "
        "또한 Ours는 입력 길이나 예측 길이가 커져도 대형 horizon-dependent head를 사용하지 않기 때문에, horizon 증가에 따른 파라미터 증가를 "
        "효과적으로 억제할 수 있다.",
    )
    add_paragraph(
        doc,
        "학습은 예측값과 실제값 사이의 평균제곱오차를 최소화하는 방식으로 수행한다.",
    )
    add_equation(doc, "L = (1/H) || y - y_hat ||_2^2")
    add_paragraph(
        doc,
        "요약하면 Ours는 미래를 직접 찍는 모델이라기보다, 미래를 구성할 basis를 만들고 이를 조합하는 모델에 가깝다. 이러한 설계는 장기 시계열 예측에 "
        "필요한 inductive bias를 제공하면서도, 파라미터 효율성을 함께 확보하는 데 목적이 있다.",
    )

    add_heading(doc, "Ⅲ. 실험 설정")
    add_paragraph(
        doc,
        "실험은 ETT(Electricity Transformer Temperature) 벤치마크의 ETTh1 데이터를 대상으로 수행하였다. ETTh1은 전력용 변압기 운용과 "
        "관련된 다변량 시계열 데이터이며, 본 실험에서는 7개 입력 변수 중 oil temperature를 의미하는 OT를 target으로 사용하였다.",
    )
    add_paragraph(
        doc,
        "입력 길이는 96, 예측 길이는 96, 192, 336, 720으로 설정하였고, 비교 모델은 Ours, DLinear, PatchTST로 구성하였다. 결과는 3개 "
        "시드 평균의 scaled MAE, scaled RMSE, 그리고 trainable parameter 수로 정리하였다. 평가지표의 초점은 절대 최고 성능뿐 아니라 모델 "
        "크기 대비 예측 효율성에 두었다.",
    )

    add_heading(doc, "Ⅳ. 실험 결과 및 분석")
    add_paragraph(
        doc,
        "표 1은 horizon별 평균 예측 오차를, 표 2는 동일 조건에서의 모델 파라미터 수를 보여준다. 두 표를 함께 보면 Ours의 성능과 모델 규모 사이의 "
        "절충 관계를 보다 명확하게 해석할 수 있다.",
    )

    add_table_caption(doc, "표 1. ETTh1 3-seed 평균 예측 성능 비교")
    table1 = doc.add_table(rows=1 + len(ACCURACY_ROWS), cols=7)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.style = "Table Grid"
    fill_table(
        table1,
        [["Pred Len", "DLinear MAE", "DLinear RMSE", "PatchTST MAE", "PatchTST RMSE", "Ours MAE", "Ours RMSE"]]
        + ACCURACY_ROWS,
    )
    set_table_borders(table1)

    add_table_caption(doc, "표 2. ETTh1 horizon별 파라미터 수 비교")
    table2 = doc.add_table(rows=1 + len(PARAM_ROWS), cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = "Table Grid"
    fill_table(
        table2,
        [["Pred Len", "DLinear Params", "PatchTST Params", "Ours Params"]] + PARAM_ROWS,
    )
    set_table_borders(table2)

    add_paragraph(
        doc,
        "실험 결과, PatchTST는 ETTh1에서 scaled MAE 기준 최고 정확도를 보였지만, 그 대가로 훨씬 큰 모델 크기를 요구한다. 반면 Ours는 모든 "
        "horizon에서 11,149개의 파라미터를 유지했고, 720-step에서는 DLinear보다 약 12.5배 적은 파라미터로 scaled MAE를 0.3221에서 "
        "0.2910으로 개선하였다. 또한 336-step과 720-step에서는 PatchTST 대비 scaled MAE 차이를 각각 1.40%와 1.06%로 제한하여, "
        "2% 정확도 허용오차 내의 parameter-efficient 대안임을 확인하였다.",
    )

    add_heading(doc, "Ⅴ. 결 론")
    add_paragraph(
        doc,
        "본 논문에서는 장기 시계열 예측을 위한 파라미터 효율적 구조 합성 모델 Ours를 제안하였다. Ours는 구조적 basis 조합을 통해 horizon 전체를 "
        "합성함으로써, 예측 길이가 증가해도 모델 크기를 일정하게 유지한다. ETTh1 실험 결과는 PatchTST가 최고 정확도를 보이는 가운데, Ours가 long horizon에서 "
        "2% 정확도 허용오차 내의 훨씬 작은 대안이 될 수 있음을 보여준다. 향후에는 추가 데이터셋 검증과 latency·memory 분석을 통해 "
        "적용 범위를 더 확장할 계획이다.",
    )

    add_heading(doc, "참 고 문 헌")
    references = [
        "[1] H. Zhou et al., Informer: Beyond Efficient Transformer for Long Sequence Time-Series Forecasting, AAAI, 2021.",
        "[2] H. Wu et al., Autoformer: Decomposition Transformers with Auto-Correlation for Long-Term Series Forecasting, NeurIPS, 2021.",
        "[3] T. Zhou et al., FEDformer: Frequency Enhanced Decomposed Transformer for Long-term Series Forecasting, ICML, 2022.",
        "[4] Y. Nie et al., A Time Series is Worth 64 Words: Long-term Forecasting with Transformers, ICLR, 2023.",
        "[5] H. Wu et al., TimesNet: Temporal 2D-Variation Modeling for General Time Series Analysis, ICLR, 2023.",
        "[6] A. Zeng et al., Are Transformers Effective for Time Series Forecasting?, AAAI, 2023.",
        "[7] B. Oreshkin et al., N-BEATS: Neural Basis Expansion Analysis for Interpretable Time Series Forecasting, ICLR, 2020.",
        "[8] Y. Liu et al., iTransformer: Inverted Transformers Are Effective for Time Series Forecasting, ICLR, 2024.",
    ]
    for ref in references:
        add_paragraph(doc, ref, size=9, space_after=1, first_line_indent_cm=None)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUT)
    print(f"OUTPUT={OUT.resolve()}")


if __name__ == "__main__":
    main()
