from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = Path("Ours_국내학회_확장초안_도표포함.docx")


SETTINGS_ROWS = [
    ["데이터셋", "ETTh1"],
    ["입력 길이", "96"],
    ["예측 길이", "96, 192, 336, 720"],
    ["비교 모델", "Ours, DLinear, PatchTST"],
    ["평가지표", "scaled MAE, scaled RMSE, trainable parameters"],
    ["실험 목적", "정확도-효율성 trade-off 검증"],
]

ACCURACY_ROWS = [
    ["96", "0.2310", "0.3031", "0.2119", "0.2794", "0.2239", "0.2900"],
    ["192", "0.2514", "0.3275", "0.2396", "0.3133", "0.2489", "0.3201"],
    ["336", "0.2655", "0.3393", "0.2553", "0.3264", "0.2589", "0.3283"],
    ["720", "0.3221", "0.3987", "0.2879", "0.3673", "0.2910", "0.3667"],
]

PARAM_ROWS = [
    ["96", "18.6K", "6903.9K", "11.1K"],
    ["192", "37.2K", "7493.8K", "11.1K"],
    ["336", "65.2K", "8378.7K", "11.1K"],
    ["720", "139.7K", "10738.4K", "11.1K"],
]

ABLATION_HEADER = ["Variant", "설명", "MAE", "비고"]
ABLATION_ROWS = [
    ["Ours", "기본 모델", "[결과 입력]", "최종 모델"],
    ["Ours-no-router", "router 제거", "[결과 입력]", "ablation"],
    ["Ours-fixed-bank", "adaptive bank 제거", "[결과 입력]", "ablation"],
    ["Ours-direct-head", "direct head 대체", "[결과 입력]", "ablation"],
]


def set_run_font(run, *, size: float, bold: bool = False, ascii_font: str = "Times New Roman", east_asia_font: str = "바탕") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def set_cell_text(cell, text: str, *, size: float = 9, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(
    doc: Document,
    text: str,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size: float = 10,
    bold: bool = False,
    space_before=0,
    space_after=0,
    first_line_indent_cm: float | None = 0.6,
):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if first_line_indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc: Document, text: str) -> None:
    add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=11,
        bold=True,
        space_before=8,
        space_after=4,
        first_line_indent_cm=None,
    )


def add_equation(doc: Document, text: str) -> None:
    add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10,
        space_before=2,
        space_after=2,
        first_line_indent_cm=None,
    )


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def add_table_caption(doc: Document, text: str) -> None:
    add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=9.5,
        space_before=4,
        space_after=2,
        first_line_indent_cm=None,
    )


def add_data_table(doc: Document, rows: list[list[str]], *, caption: str) -> None:
    add_table_caption(doc, caption)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            set_cell_text(table.rows[r_idx].cells[c_idx], text, bold=(r_idx == 0))


def add_placeholder_figure(doc: Document, *, title: str, note: str, height_lines: int = 6) -> None:
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.style = "Table Grid"
    set_table_borders(box)
    cell = box.cell(0, 0)
    cell.width = Cm(15.5)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[{title}]")
    set_run_font(run, size=10, bold=True)
    for _ in range(height_lines):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" ")
        set_run_font(run, size=9)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(note)
    set_run_font(run, size=9)


def configure_page(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def build_document() -> Document:
    doc = Document()
    configure_page(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ours: 장기 시계열 예측을 위한 파라미터 효율적 구조 합성 모델")
    set_run_font(run, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ours: A Parameter-Efficient Structural Synthesis Model for Long-Term Time Series Forecasting")
    set_run_font(run, size=11, bold=True)

    for text, size in [
        ("[저자명1], [저자명2]*", 11),
        ("[소속기관명], *[공동소속기관명]", 10),
        ("author1@domain.ac.kr, *author2@domain.ac.kr", 10),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=size)

    add_heading(doc, "요 약")
    add_paragraph(
        doc,
        "장기 시계열 예측에서는 horizon이 길어질수록 예측 오차뿐 아니라 모델 규모와 계산 비용도 함께 증가한다. "
        "본 논문은 미래 구간 전체를 큰 출력층으로 직접 회귀하는 대신, 과거 시계열의 잠재 요약으로부터 구조적 basis를 생성하고 "
        "이를 조합하여 horizon 전체를 합성하는 Ours를 제안한다. ETTh1 데이터셋에서 입력 길이 96, 예측 길이 96, 192, 336, 720 조건으로 "
        "Ours를 DLinear와 PatchTST와 비교한 결과, Ours는 모든 horizon에서 11,149개의 파라미터를 유지하면서 DLinear보다 낮은 "
        "scaled MAE를 기록하였다. 또한 336-step과 720-step에서는 PatchTST의 최고 scaled MAE 대비 2% 이내의 오차를 유지하면서 "
        "훨씬 작은 모델 규모를 보였다. 이 결과는 Ours가 장기 시계열 예측에서 정확도와 파라미터 효율성 사이의 실용적인 절충점을 제공함을 보여준다.",
        size=10,
    )

    add_heading(doc, "Ⅰ. 서 론")
    add_paragraph(
        doc,
        "장기 시계열 예측은 전력 수요, 설비 상태, 교통 흐름 등 다양한 응용에서 중요한 문제이며, 예측 horizon이 길어질수록 오차 누적과 "
        "장기 의존성 모델링의 난도가 함께 증가한다. 특히 미래 구간이 길어질수록 출력층 규모와 계산 비용까지 커지기 때문에, 정확도뿐 아니라 "
        "구조적 효율성도 함께 고려한 모델 설계가 필요하다."
    )
    add_paragraph(
        doc,
        "최근 Informer[1], Autoformer[2], FEDformer[3]와 같은 Transformer 계열 모델과 PatchTST[4], TimesNet[5], "
        "iTransformer[8] 등이 장기 예측 성능을 개선해 왔다. 반면 Are Transformers Effective for Time Series Forecasting?[6]은 "
        "단순한 linear 계열 모델도 강한 기준선이 될 수 있음을 보였으며, 이는 backbone의 복잡성만으로는 장기 예측 성능을 충분히 설명하기 어렵다는 점을 시사한다."
    )
    add_paragraph(
        doc,
        "따라서 본 연구는 '훨씬 작은 모델로도 경쟁력 있는 장기 예측 성능을 낼 수 있는가'라는 질문에 초점을 맞춘다. 제안 모델 Ours는 "
        "예측 horizon 전체를 구조적 basis 조합으로 생성하는 방식으로 설계되며, 이를 통해 direct multi-step head가 갖는 파라미터 증가 문제를 "
        "완화하고자 한다."
    )

    add_heading(doc, "Ⅱ. 관련 연구 및 문제 정의")
    add_paragraph(
        doc,
        "기존 장기 시계열 예측 모델은 크게 Transformer 기반 접근과 linear 기반 접근으로 나눌 수 있다. Transformer 계열은 복잡한 attention "
        "구조를 통해 장기 의존성을 포착하지만, 큰 표현 차원과 다층 encoder-decoder 구조로 인해 모델 규모가 커지는 경향이 있다."
    )
    add_paragraph(
        doc,
        "반면 linear 계열 모델은 구조가 단순하고 학습이 안정적이지만, 긴 horizon에서 복합적인 추세·주기·과도 패턴을 동시에 표현하는 데 한계가 있을 수 있다. "
        "본 논문은 이 둘의 중간 지점에서, 장기 예측에 필요한 구조적 inductive bias를 유지하면서도 모델 규모를 작게 가져갈 수 있는 설계를 지향한다."
    )
    add_paragraph(
        doc,
        "문제 정의 측면에서, 입력 시계열 X ∈ R^(L×C)가 주어졌을 때 목표는 미래 horizon H에 대한 target channel의 예측값 y_hat ∈ R^H를 "
        "생성하는 것이다. Ours는 이 문제를 시점별 직접 회귀가 아니라 구조적 basis 합성 문제로 재정의한다."
    )

    add_heading(doc, "Ⅲ. 제안 방법")
    add_paragraph(
        doc,
        "입력 다변량 시계열을 X ∈ R^(L×C), 예측 horizon을 H라 하자. Ours는 입력 시계열을 causal encoder에 통과시켜 시간 순서를 보존하는 "
        "잠재 표현과 요약 벡터를 얻는다."
    )
    add_equation(doc, "E = f_enc(X),      z = Pool(E)")
    add_paragraph(
        doc,
        "그 다음 요약 벡터 z로부터 미래 구간을 설명하는 구조적 basis와 각 basis의 계수를 예측한다. 본 연구에서 사용하는 basis는 추세형(trend), "
        "주기형(seasonal), 과도형(transient) 성분으로 구성되며, Ours의 핵심은 horizon의 각 시점을 독립적으로 예측하지 않고 horizon 전체를 "
        "하나의 구조적 조합 문제로 다룬다는 점이다."
    )
    add_equation(doc, "B = g_basis(z),      a = g_coef(z),      y_hat = Σ_{k=1}^{K} a_k B_k")
    add_paragraph(
        doc,
        "trend basis는 장기 증가·감소 경향을, seasonal basis는 반복 주기를, transient basis는 단기 감쇠 패턴을 담당한다. "
        "또한 Ours는 입력 길이나 예측 길이가 커져도 대형 horizon-dependent head를 사용하지 않기 때문에, horizon 증가에 따른 파라미터 증가를 "
        "효과적으로 억제할 수 있다."
    )
    add_paragraph(doc, "학습은 예측값과 실제값 사이의 평균제곱오차를 최소화하는 방식으로 수행한다.")
    add_equation(doc, "L = (1/H) || y - y_hat ||_2^2")
    add_paragraph(
        doc,
        "요약하면 Ours는 미래를 직접 찍는 모델이라기보다, 미래를 구성할 basis를 만들고 이를 조합하는 모델에 가깝다. 이러한 설계는 장기 시계열 예측에 "
        "필요한 inductive bias를 제공하면서도, 파라미터 효율성을 함께 확보하는 데 목적이 있다."
    )

    add_table_caption(doc, "그림 1. Ours 전체 구조도 삽입 위치")
    add_placeholder_figure(
        doc,
        title="그림 1 삽입 위치",
        note="여기에 encoder, summary, basis generator, coefficient head, synthesis 흐름을 나타내는 구조도를 삽입",
        height_lines=5,
    )

    add_heading(doc, "Ⅳ. 실험 설정")
    add_paragraph(
        doc,
        "실험은 ETT(Electricity Transformer Temperature) 벤치마크의 ETTh1 데이터를 대상으로 수행하였다. ETTh1은 전력용 변압기 운용과 "
        "관련된 다변량 시계열 데이터이며, 본 실험에서는 7개 입력 변수 중 oil temperature를 의미하는 OT를 target으로 사용하였다."
    )
    add_paragraph(
        doc,
        "입력 길이는 96, 예측 길이는 96, 192, 336, 720으로 설정하였고, 비교 모델은 Ours, DLinear, PatchTST로 구성하였다. 결과는 3개 "
        "시드 평균의 scaled MAE, scaled RMSE, 그리고 trainable parameter 수로 정리하였다. 평가지표의 초점은 절대 최고 성능뿐 아니라 모델 "
        "크기 대비 예측 효율성에 두었다."
    )
    add_data_table(doc, [["항목", "설정"]] + SETTINGS_ROWS, caption="표 1. 실험 설정 요약")

    add_heading(doc, "Ⅴ. 실험 결과 및 논의")
    add_paragraph(
        doc,
        "표 2는 horizon별 평균 예측 오차를, 표 3은 동일 조건에서의 모델 파라미터 수를 보여준다. 두 표를 함께 보면 Ours의 성능과 모델 규모 "
        "사이의 절충 관계를 보다 명확하게 해석할 수 있다."
    )
    add_data_table(
        doc,
        [["Pred Len", "DLinear MAE", "DLinear RMSE", "PatchTST MAE", "PatchTST RMSE", "Ours MAE", "Ours RMSE"]]
        + ACCURACY_ROWS,
        caption="표 2. ETTh1 3-seed 평균 예측 성능 비교",
    )
    add_data_table(
        doc,
        [["Pred Len", "DLinear Params", "PatchTST Params", "Ours Params"]] + PARAM_ROWS,
        caption="표 3. ETTh1 horizon별 파라미터 수 비교",
    )
    add_paragraph(
        doc,
        "실험 결과, PatchTST는 ETTh1에서 scaled MAE 기준 최고 정확도를 보였지만, 그 대가로 훨씬 큰 모델 크기를 요구한다. 반면 Ours는 모든 "
        "horizon에서 11,149개의 파라미터를 유지했고, 720-step에서는 DLinear보다 약 12.5배 적은 파라미터로 scaled MAE를 0.3221에서 "
        "0.2910으로 개선하였다."
    )
    add_paragraph(
        doc,
        "또한 336-step과 720-step에서는 PatchTST 대비 scaled MAE 차이를 각각 1.40%와 1.06%로 제한하여, 2% 정확도 허용오차 내의 "
        "parameter-efficient 대안임을 확인하였다. 이는 Ours가 '최고 정확도'를 무조건 목표로 하기보다, 실사용 관점에서 충분히 강한 성능과 작은 모델 규모를 동시에 추구하는 접근임을 보여준다."
    )

    add_table_caption(doc, "그림 2. 정확도-효율성 trade-off 그래프 삽입 위치")
    add_placeholder_figure(
        doc,
        title="그림 2 삽입 위치",
        note="여기에 MAE-parameter 혹은 MAE-latency scatter plot을 삽입",
        height_lines=5,
    )

    add_paragraph(
        doc,
        "추가로, 실제 논문화 단계에서는 예측 곡선 시각화와 ablation 결과가 함께 제시될 필요가 있다. 예측 곡선은 모델이 단순 평균 회귀에 그치지 않고 "
        "장기 추세와 반복 패턴을 어느 정도 따라가는지 직관적으로 보여주며, ablation 표는 router, bank, direct-head 대체 등 각 설계 요소의 기여를 "
        "정량적으로 설명하는 데 유용하다."
    )

    add_table_caption(doc, "그림 3. 예측 사례 시각화 삽입 위치")
    add_placeholder_figure(
        doc,
        title="그림 3 삽입 위치",
        note="여기에 실제값과 Ours/DLinear/PatchTST 예측 곡선 비교 그림을 삽입",
        height_lines=4,
    )

    add_data_table(
        doc,
        [ABLATION_HEADER] + ABLATION_ROWS,
        caption="표 4. Ours ablation 결과 삽입용 표",
    )

    add_heading(doc, "Ⅵ. 결 론")
    add_paragraph(
        doc,
        "본 논문에서는 장기 시계열 예측을 위한 파라미터 효율적 구조 합성 모델 Ours를 제안하였다. Ours는 구조적 basis 조합을 통해 horizon 전체를 "
        "합성함으로써, 예측 길이가 증가해도 모델 크기를 일정하게 유지한다. ETTh1 실험 결과는 PatchTST가 최고 정확도를 보이는 가운데, Ours가 long horizon에서 "
        "2% 정확도 허용오차 내의 훨씬 작은 대안이 될 수 있음을 보여준다."
    )
    add_paragraph(
        doc,
        "향후에는 추가 데이터셋 검증, latency·memory 분석, 그리고 variant별 ablation 실험을 통해 적용 범위를 확장할 계획이다. 또한 최종 제출본에서는 "
        "본 초안에 표시한 그림 및 보조 표를 실제 실험 결과로 채워 넣어, Ours의 구조적 장점과 accuracy-efficiency trade-off를 보다 설득력 있게 "
        "제시할 수 있을 것으로 기대한다."
    )

    add_heading(doc, "참 고 문 헌")
    references = [
        "[1] H. Zhou et al., Informer: Beyond Efficient Transformer for Long Sequence Time-Series Forecasting, AAAI, 2021.",
        "[2] H. Wu et al., Autoformer: Decomposition Transformers with Auto-Correlation for Long-Term Series Forecasting, NeurIPS, 2021.",
        "[3] T. Zhou et al., FEDformer: Frequency Enhanced Decomposed Transformer for Long-term Series Forecasting, ICML, 2022.",
        "[4] Y. Nie et al., A Time Series is Worth 64 Words: Long-term Forecasting with Transformers, ICLR, 2023.",
        "[5] H. Wu et al., TimesNet: Temporal 2D-Variation Modeling for General Time Series Analysis, ICLR, 2023.",
        "[6] A. Zeng et al., Are Transformers Effective for Time Series Forecasting?, AAAI, 2023.",
        "[7] B. Oreshkin et al., N-BEATS: Neural Basis Expansion Analysis for Interpretable Time Series Forecasting, ICLR, 2020.",
        "[8] Y. Liu et al., iTransformer: Inverted Transformers Are Effective for Time Series Forecasting, ICLR, 2024.",
    ]
    for ref in references:
        add_paragraph(doc, ref, size=9, space_after=1, first_line_indent_cm=None)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUT)
    print(f"OUTPUT={OUT.resolve()}")


if __name__ == "__main__":
    main()
